from __future__ import annotations

'''Safely stage the authoritative Chapters I-III DOCX from its Markdown source.

The reference Word package supplies the cover, author table, chapter/page-break
anchors, References, hyperlinks, styles, and section geometry. Pandoc renders
the current Markdown with that reference, after which only the cover title and
the bodies of Chapters I-III are transferred into a clone. The reference file
is never overwritten by this script.
'''

import argparse
import copy
import hashlib
import json
import re
import subprocess
import tempfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


EXPECTED_REFERENCE_SHA256 = (
    'D920524D78C871F7E503675C423C6F82573C3BE5E53D7587E7D4005BF11D9FE3'
)
PANDOC_READER = (
    'markdown+pipe_tables+tex_math_dollars+tex_math_single_backslash'
)
ANCHORS = (
    'I. Introduction',
    'II. RELATED WORK',
    'III. METHODOLOGY',
    'References',
)
CHAPTERS = ANCHORS[:3]

W_P = qn('w:p')
W_TBL = qn('w:tbl')
W_SECT_PR = qn('w:sectPr')
W_P_PR = qn('w:pPr')
W_R_PR = qn('w:rPr')
W_T = qn('w:t')
W_BR = qn('w:br')
W_HYPERLINK = qn('w:hyperlink')
W_BOOKMARK_START = qn('w:bookmarkStart')
W_BOOKMARK_END = qn('w:bookmarkEnd')
W_NUM_PR = qn('w:numPr')
R_ID = qn('r:id')
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'


class SyncError(RuntimeError):
    pass


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open('rb') as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b''):
            digest.update(chunk)
    return digest.hexdigest().upper()


def normalize(value: str) -> str:
    return re.sub(r'\s+', ' ', value or '').strip()


def node_text(node: etree._Element) -> str:
    pieces: list[str] = []
    for descendant in node.iter():
        if descendant.tag == W_T:
            pieces.append(descendant.text or '')
        elif descendant.tag == W_BR:
            pieces.append(' ')
    return normalize(''.join(pieces))


def direct_children(document: Any) -> list[etree._Element]:
    return list(document.element.body.iterchildren())


def is_page_break(node: etree._Element) -> bool:
    if node.tag != W_P or node_text(node):
        return False
    return any(
        child.get(qn('w:type'), '') == 'page'
        for child in node.iter(W_BR)
    )


def canonical_fingerprint(nodes: list[etree._Element]) -> str:
    digest = hashlib.sha256()
    for node in nodes:
        payload = etree.tostring(node, method='c14n', with_comments=True)
        digest.update(len(payload).to_bytes(8, 'big'))
        digest.update(payload)
    return digest.hexdigest().upper()


def find_anchor(
    document: Any,
    text: str,
) -> tuple[int, etree._Element]:
    matches = [
        (index, node)
        for index, node in enumerate(direct_children(document))
        if node.tag == W_P and node_text(node) == text
    ]
    if len(matches) != 1:
        raise SyncError(
            f'Expected exactly one direct-body anchor {text!r}; '
            f'found {len(matches)}'
        )
    return matches[0]


def anchor_map(document: Any) -> dict[str, etree._Element]:
    found = {name: find_anchor(document, name)[1] for name in ANCHORS}
    positions = [direct_children(document).index(found[name]) for name in ANCHORS]
    if positions != sorted(positions):
        raise SyncError(f'Heading anchors are out of order: {positions}')
    return found


def page_break_before(
    document: Any,
    next_anchor: etree._Element,
) -> etree._Element:
    children = direct_children(document)
    next_index = children.index(next_anchor)
    candidates = [
        node
        for node in children[:next_index]
        if is_page_break(node)
    ]
    if not candidates:
        raise SyncError(
            f'No page-break paragraph precedes {node_text(next_anchor)!r}'
        )
    candidate = candidates[-1]
    prior_anchor_positions = [
        index
        for index, node in enumerate(children)
        if node.tag == W_P and node_text(node) in ANCHORS
    ]
    previous_anchor_index = max(
        index for index in prior_anchor_positions if index < next_index
    )
    candidate_index = children.index(candidate)
    if not previous_anchor_index < candidate_index < next_index:
        raise SyncError(
            f'Page-break placement is ambiguous before {node_text(next_anchor)!r}'
        )
    return candidate


def reference_regions(document: Any) -> dict[str, Any]:
    anchors = anchor_map(document)
    children = direct_children(document)
    title_candidates = [
        node
        for node in children[:children.index(anchors[ANCHORS[0]])]
        if node.tag == W_P and node_text(node)
    ]
    if not title_candidates:
        raise SyncError('The reference cover has no title paragraph')
    title = title_candidates[0]
    cover_nodes = children[:children.index(anchors[ANCHORS[0]])]
    cover_without_title = [node for node in cover_nodes if node is not title]
    author_tables = [node for node in cover_nodes if node.tag == W_TBL]
    if len(author_tables) != 1:
        raise SyncError(
            f'Expected one cover author table; found {len(author_tables)}'
        )
    refs_index = children.index(anchors['References'])
    page_breaks = [
        page_break_before(document, anchors[ANCHORS[index]])
        for index in range(1, len(ANCHORS))
    ]
    if len({id(node) for node in page_breaks}) != 3:
        raise SyncError('Expected three distinct chapter page breaks')
    return {
        'anchors': anchors,
        'title': title,
        'cover_without_title': cover_without_title,
        'author_table': author_tables[0],
        'references': children[refs_index:],
        'page_breaks': page_breaks,
    }


def markdown_title(source: Path) -> str:
    first = next(
        (line.strip() for line in source.read_text(encoding='utf-8').splitlines()
         if line.strip()),
        '',
    )
    match = re.fullmatch(r'__\*(.+)\*__', first)
    if not match:
        raise SyncError(
            'The first non-empty Markdown line must be the bold-italic title'
        )
    return match.group(1).strip()


def first_run_properties(paragraph: etree._Element) -> etree._Element | None:
    for descendant in paragraph.iter(W_R_PR):
        return copy.deepcopy(descendant)
    return None


def replace_paragraph_text(paragraph: etree._Element, value: str) -> None:
    p_pr = paragraph.find(W_P_PR)
    run_properties = first_run_properties(paragraph)
    for child in list(paragraph):
        if child is not p_pr:
            paragraph.remove(child)
    run = OxmlElement('w:r')
    if run_properties is not None:
        run.append(run_properties)
    text = OxmlElement('w:t')
    text.set(XML_SPACE, 'preserve')
    text.text = value
    run.append(text)
    paragraph.append(run)


def source_inner_nodes(
    generated: Any,
    chapter: str,
    next_anchor: str,
) -> list[etree._Element]:
    start_index, _ = find_anchor(generated, chapter)
    end_index, _ = find_anchor(generated, next_anchor)
    if start_index >= end_index:
        raise SyncError(f'Generated chapter order is invalid for {chapter!r}')
    selected: list[etree._Element] = []
    for node in direct_children(generated)[start_index + 1:end_index]:
        if node.tag in {W_BOOKMARK_START, W_BOOKMARK_END}:
            continue
        if node.tag not in {W_P, W_TBL}:
            raise SyncError(
                f'Unsupported generated body node in {chapter!r}: {node.tag}'
            )
        if is_page_break(node):
            continue
        if any(descendant.get(R_ID) for descendant in node.iter()):
            raise SyncError(
                f'Generated {chapter!r} contains relationship-bound content'
            )
        cloned = copy.deepcopy(node)
        for tag in (W_BOOKMARK_START, W_BOOKMARK_END):
            for bookmark in list(cloned.iter(tag)):
                parent = bookmark.getparent()
                if parent is not None:
                    parent.remove(bookmark)
        selected.append(cloned)
    if not selected:
        raise SyncError(f'Generated chapter {chapter!r} is empty')
    return selected


def replace_part_children(
    destination: etree._Element,
    source: etree._Element,
) -> None:
    for key in list(destination.attrib):
        del destination.attrib[key]
    for key, value in source.attrib.items():
        destination.set(key, value)
    for child in list(destination):
        destination.remove(child)
    for child in source:
        destination.append(copy.deepcopy(child))


def set_content_table_geometry(table: Any, widths: list[int]) -> None:
    if len(table.columns) != len(widths):
        raise SyncError(
            f'Table column mismatch: {len(table.columns)} vs {len(widths)}'
        )
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_style = tbl_pr.find(qn('w:tblStyle'))
    if tbl_style is None:
        tbl_style = OxmlElement('w:tblStyle')
        tbl_pr.insert(0, tbl_style)
    tbl_style.set(qn('w:val'), 'TableGrid')
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_w.set(qn('w:w'), str(total))
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:type'), 'dxa')
    # Match the default 120-DXA start cell margin so the visible border aligns
    # with the surrounding body text under the packaged geometry audit.
    tbl_ind.set(qn('w:w'), '120')
    layout = tbl_pr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tbl_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement('w:gridCol')
        column.set(qn('w:w'), str(width))
        grid.append(column)

    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        for height in list(tr_pr.findall(qn('w:trHeight'))):
            tr_pr.remove(height)
        if tr_pr.find(qn('w:cantSplit')) is None:
            tr_pr.append(OxmlElement('w:cantSplit'))
        if row_index == 0 and tr_pr.find(qn('w:tblHeader')) is None:
            tr_pr.append(OxmlElement('w:tblHeader'))
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:type'), 'dxa')
            tc_w.set(qn('w:w'), str(width))
            for paragraph in cell.paragraphs:
                p_pr = paragraph._p.get_or_add_pPr()
                p_style = p_pr.find(qn('w:pStyle'))
                if p_style is None:
                    p_style = OxmlElement('w:pStyle')
                    p_pr.insert(0, p_style)
                p_style.set(qn('w:val'), 'Normal')
                spacing = p_pr.find(qn('w:spacing'))
                if spacing is None:
                    spacing = OxmlElement('w:spacing')
                    p_pr.append(spacing)
                spacing.set(qn('w:after'), '0')
                if row_index == 0:
                    for run in paragraph.runs:
                        r_pr = run._r.get_or_add_rPr()
                        if r_pr.find(qn('w:b')) is None:
                            r_pr.append(OxmlElement('w:b'))


def append_simple_numbering(
    numbering: etree._Element,
    abstract_id: int,
    num_id: int,
    kind: str,
    left: int,
) -> None:
    abstract = OxmlElement('w:abstractNum')
    abstract.set(qn('w:abstractNumId'), str(abstract_id))
    nsid = OxmlElement('w:nsid')
    nsid.set(qn('w:val'), f'{abstract_id:08X}')
    abstract.append(nsid)
    multilevel = OxmlElement('w:multiLevelType')
    multilevel.set(qn('w:val'), 'singleLevel')
    abstract.append(multilevel)
    level = OxmlElement('w:lvl')
    level.set(qn('w:ilvl'), '0')
    start = OxmlElement('w:start')
    start.set(qn('w:val'), '1')
    level.append(start)
    num_fmt = OxmlElement('w:numFmt')
    num_fmt.set(qn('w:val'), 'bullet' if kind == 'bullet' else 'decimal')
    level.append(num_fmt)
    level_text = OxmlElement('w:lvlText')
    level_text.set(qn('w:val'), '•' if kind == 'bullet' else '%1.')
    if kind == 'bullet':
        level_text.set(qn('w:val'), '\u2022')
    level.append(level_text)
    justification = OxmlElement('w:lvlJc')
    justification.set(qn('w:val'), 'left')
    level.append(justification)
    paragraph_properties = OxmlElement('w:pPr')
    indent = OxmlElement('w:ind')
    indent.set(qn('w:left'), str(left))
    indent.set(qn('w:hanging'), '360')
    paragraph_properties.append(indent)
    level.append(paragraph_properties)
    if kind == 'bullet':
        run_properties = OxmlElement('w:rPr')
        fonts = OxmlElement('w:rFonts')
        fonts.set(qn('w:ascii'), 'Arial')
        fonts.set(qn('w:hAnsi'), 'Arial')
        run_properties.append(fonts)
        level.append(run_properties)
    abstract.append(level)
    first_concrete = numbering.find(qn('w:num'))
    if first_concrete is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_concrete), abstract)

    concrete = OxmlElement('w:num')
    concrete.set(qn('w:numId'), str(num_id))
    abstract_ref = OxmlElement('w:abstractNumId')
    abstract_ref.set(qn('w:val'), str(abstract_id))
    concrete.append(abstract_ref)
    numbering.append(concrete)


def normalize_list_numbering(document: Any) -> dict[str, int]:
    paragraphs: list[tuple[etree._Element, str]] = []
    ordered_groups: list[str] = []
    for paragraph in document.element.body.iter(W_P):
        p_pr = paragraph.find(W_P_PR)
        num_pr = p_pr.find(W_NUM_PR) if p_pr is not None else None
        num_id_node = num_pr.find(qn('w:numId')) if num_pr is not None else None
        if num_id_node is None:
            continue
        old_id = num_id_node.get(qn('w:val'), '')
        paragraphs.append((paragraph, old_id))
        if old_id not in ordered_groups:
            ordered_groups.append(old_id)
    counts = Counter(old_id for _, old_id in paragraphs)
    expected_counts = [4, 1, 3, 7, 4]
    actual_counts = [counts[group] for group in ordered_groups]
    if actual_counts != expected_counts:
        raise SyncError(
            f'Unexpected generated list groups/counts: '
            f'{list(zip(ordered_groups, actual_counts))}'
        )

    numbering = document.part.numbering_part.element
    kinds = ['decimal', 'bullet', 'decimal', 'decimal', 'decimal']
    left_indents = [720, 720, 1080, 720, 720]
    mapping: dict[str, int] = {}
    for index, (old_id, kind, left) in enumerate(
        zip(ordered_groups, kinds, left_indents)
    ):
        new_id = 20 + index
        mapping[old_id] = new_id
        append_simple_numbering(numbering, new_id, new_id, kind, left)

    for paragraph, old_id in paragraphs:
        num_pr = paragraph.find(W_P_PR).find(W_NUM_PR)
        num_pr.find(qn('w:numId')).set(qn('w:val'), str(mapping[old_id]))
        level = num_pr.find(qn('w:ilvl'))
        if level is None:
            level = OxmlElement('w:ilvl')
            num_pr.insert(0, level)
        level.set(qn('w:val'), '0')
    return {old_id: mapping[old_id] for old_id in ordered_groups}


def materialize_list_markers(document: Any, numbering_ids: set[int]) -> int:
    '''Make list markers viewer-independent while preserving hanging indents.'''
    sequence: dict[int, int] = {numbering_id: 0 for numbering_id in numbering_ids}
    materialized = 0
    for paragraph in document.element.body.iter(W_P):
        p_pr = paragraph.find(W_P_PR)
        num_pr = p_pr.find(W_NUM_PR) if p_pr is not None else None
        num_id_node = num_pr.find(qn('w:numId')) if num_pr is not None else None
        if num_id_node is None:
            continue
        num_id = int(num_id_node.get(qn('w:val'), '-1'))
        if num_id not in numbering_ids:
            continue
        sequence[num_id] += 1
        prefix = '\u2022 ' if num_id == 21 else f'{sequence[num_id]}. '
        p_pr.remove(num_pr)
        indent = p_pr.find(qn('w:ind'))
        if indent is None:
            indent = OxmlElement('w:ind')
            p_pr.append(indent)
        indent.set(qn('w:left'), '720' if num_id != 22 else '1080')
        indent.set(qn('w:hanging'), '360')
        marker_run = OxmlElement('w:r')
        marker_text = OxmlElement('w:t')
        marker_text.set(XML_SPACE, 'preserve')
        marker_text.text = prefix
        marker_run.append(marker_text)
        paragraph.insert(1, marker_run)
        materialized += 1
    return materialized


def replace_chapter(
    output: Any,
    chapter_anchor: etree._Element,
    page_break: etree._Element,
    new_nodes: list[etree._Element],
) -> int:
    body = output.element.body
    children = direct_children(output)
    start = children.index(chapter_anchor) + 1
    end = children.index(page_break)
    if start > end:
        raise SyncError(f'Invalid replacement range for {node_text(chapter_anchor)}')
    old_nodes = children[start:end]
    for node in old_nodes:
        body.remove(node)
    insertion = list(body.iterchildren()).index(page_break)
    for node in new_nodes:
        body.insert(insertion, node)
        insertion += 1
    return len(old_nodes)


def count_direct_page_breaks(document: Any) -> int:
    return sum(is_page_break(node) for node in direct_children(document))


def count_hyperlinks(document: Any) -> int:
    return sum(1 for _ in document.element.iter(W_HYPERLINK))


def count_numbered_paragraphs(document: Any) -> int:
    count = 0
    for node in document.element.body.iter(W_P):
        p_pr = node.find(W_P_PR)
        if p_pr is not None and p_pr.find(W_NUM_PR) is not None:
            count += 1
    return count


def count_visible_list_markers(document: Any) -> int:
    expected = {'1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '\u2022 '}
    count = 0
    for node in document.element.body.iter(W_P):
        first_text = node.find('.//' + W_T)
        if first_text is not None and first_text.text in expected:
            count += 1
    return count


def heading_counts(document: Any) -> dict[str, int]:
    counts = {'Heading 1': 0, 'Heading 2': 0, 'Heading 3': 0}
    for paragraph in document.paragraphs:
        name = paragraph.style.name if paragraph.style is not None else ''
        if name in counts:
            counts[name] += 1
    return counts


def tracked_or_field_counts(document: Any) -> dict[str, int]:
    tags = {
        'comments': qn('w:commentRangeStart'),
        'insertions': qn('w:ins'),
        'deletions': qn('w:del'),
        'fields': qn('w:fldSimple'),
        'field_codes': qn('w:instrText'),
    }
    return {
        name: sum(1 for _ in document.element.iter(tag))
        for name, tag in tags.items()
    }


def validate_geometry(document: Any) -> dict[str, float]:
    if len(document.sections) != 1:
        raise SyncError(f'Expected one section; found {len(document.sections)}')
    section = document.sections[0]
    values = {
        'page_width_inches': section.page_width.inches,
        'page_height_inches': section.page_height.inches,
        'top_margin_inches': section.top_margin.inches,
        'right_margin_inches': section.right_margin.inches,
        'bottom_margin_inches': section.bottom_margin.inches,
        'left_margin_inches': section.left_margin.inches,
    }
    expected = {
        'page_width_inches': 8.5,
        'page_height_inches': 11.0,
        'top_margin_inches': 1.0,
        'right_margin_inches': 1.0,
        'bottom_margin_inches': 1.0,
        'left_margin_inches': 1.0,
    }
    for key, wanted in expected.items():
        if abs(values[key] - wanted) > 0.01:
            raise SyncError(f'Unexpected section geometry {key}={values[key]}')
    return values


def validate_output(
    document: Any,
    expected_title: str,
    before: dict[str, str],
) -> dict[str, Any]:
    regions = reference_regions(document)
    if node_text(regions['title']) != expected_title:
        raise SyncError('The synchronized cover title does not match Markdown')
    preserved = {
        'cover_without_title': canonical_fingerprint(
            regions['cover_without_title']
        ),
        'cover_author_table': canonical_fingerprint([regions['author_table']]),
        'references': canonical_fingerprint(regions['references']),
        'page_breaks': canonical_fingerprint(regions['page_breaks']),
    }
    if preserved != before:
        raise SyncError(
            'A protected cover, page-break, author-table, or References region changed'
        )
    headings = heading_counts(document)
    if headings != {'Heading 1': 4, 'Heading 2': 15, 'Heading 3': 20}:
        raise SyncError(f'Unexpected heading counts: {headings}')
    if count_direct_page_breaks(document) != 3:
        raise SyncError('The synchronized document must retain three page breaks')
    if count_hyperlinks(document) != 13:
        raise SyncError('The synchronized document must retain 13 hyperlinks')
    table_shapes = [
        [len(table.rows), len(table.columns)]
        for table in document.tables
    ]
    if len(table_shapes) != 3 or table_shapes[1:] != [[8, 2], [4, 3]]:
        raise SyncError(f'Unexpected table structure: {table_shapes}')
    numbered = count_numbered_paragraphs(document)
    visible_markers = count_visible_list_markers(document)
    if numbered != 0 or visible_markers != 19:
        raise SyncError(
            'Expected 19 viewer-independent list markers and no automatic '
            f'numbering; found visible={visible_markers}, automatic={numbered}'
        )
    special = tracked_or_field_counts(document)
    if any(special.values()):
        raise SyncError(f'Unexpected tracked changes/comments/fields: {special}')
    return {
        'heading_counts': headings,
        'direct_page_breaks': 3,
        'hyperlinks': 13,
        'table_shapes': table_shapes,
        'automatic_numbered_paragraphs': numbered,
        'visible_list_paragraphs': visible_markers,
        'special_content_counts': special,
        'section_geometry': validate_geometry(document),
        'preserved_fingerprints': preserved,
    }


def build(args: argparse.Namespace) -> dict[str, Any]:
    reference = Path(args.reference).resolve()
    source = Path(args.source).resolve()
    output = Path(args.output).resolve()
    report = Path(args.report).resolve()
    pandoc = Path(args.pandoc).resolve()
    for path in (reference, source, pandoc):
        if not path.exists():
            raise SyncError(f'Required input does not exist: {path}')
    if output == reference:
        raise SyncError('Output must be a staging path, not the reference DOCX')
    if output.exists():
        raise SyncError(f'Refusing to overwrite existing staging output: {output}')
    reference_hash = sha256_file(reference)
    if reference_hash != EXPECTED_REFERENCE_SHA256:
        raise SyncError(
            'Reference DOCX hash differs from the audited authority; '
            f'expected {EXPECTED_REFERENCE_SHA256}, found {reference_hash}'
        )

    original = Document(reference)
    regions = reference_regions(original)
    protected_before = {
        'cover_without_title': canonical_fingerprint(
            regions['cover_without_title']
        ),
        'cover_author_table': canonical_fingerprint([regions['author_table']]),
        'references': canonical_fingerprint(regions['references']),
        'page_breaks': canonical_fingerprint(regions['page_breaks']),
    }
    title = markdown_title(source)

    with tempfile.TemporaryDirectory(prefix='thesis-docx-sync-') as temp_dir:
        generated_path = Path(temp_dir) / 'pandoc_generated.docx'
        command = [
            str(pandoc),
            str(source),
            '-f',
            PANDOC_READER,
            '-t',
            'docx',
            '--reference-doc',
            str(reference),
            '-o',
            str(generated_path),
        ]
        completed = subprocess.run(
            command,
            check=False,
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='replace',
        )
        if completed.returncode != 0 or not generated_path.exists():
            raise SyncError(
                'Pandoc generation failed: '
                + (completed.stderr.strip() or completed.stdout.strip())
            )
        generated = Document(generated_path)
        generated_anchors = anchor_map(generated)

        output_doc = Document(reference)
        output_regions = reference_regions(output_doc)
        replace_paragraph_text(output_regions['title'], title)

        copied_counts: dict[str, int] = {}
        removed_counts: dict[str, int] = {}
        for index, chapter in enumerate(CHAPTERS):
            next_anchor = ANCHORS[index + 1]
            nodes = source_inner_nodes(generated, chapter, next_anchor)
            copied_counts[chapter] = len(nodes)
            target_anchor = output_regions['anchors'][chapter]
            target_next = output_regions['anchors'][next_anchor]
            page_break = page_break_before(output_doc, target_next)
            removed_counts[chapter] = replace_chapter(
                output_doc,
                target_anchor,
                page_break,
                nodes,
            )

        replace_part_children(
            output_doc.part.numbering_part.element,
            generated.part.numbering_part.element,
        )
        replace_part_children(
            output_doc.part._styles_part.element,
            generated.part._styles_part.element,
        )
        numbering_map = normalize_list_numbering(output_doc)
        materialized_list_count = materialize_list_markers(
            output_doc,
            set(numbering_map.values()),
        )
        content_tables = output_doc.tables[1:]
        if len(content_tables) != 2:
            raise SyncError(
                f'Expected two content tables; found {len(content_tables)}'
            )
        set_content_table_geometry(content_tables[0], [6624, 2736])
        set_content_table_geometry(content_tables[1], [2592, 4896, 1872])

        output.parent.mkdir(parents=True, exist_ok=True)
        output_doc.save(output)

    reopened = Document(output)
    qa = validate_output(reopened, title, protected_before)
    receipt = {
        'schema_version': 1,
        'created_utc': datetime.now(timezone.utc).isoformat(),
        'reference': str(reference),
        'source': str(source),
        'output': str(output),
        'reference_sha256': reference_hash,
        'source_sha256': sha256_file(source),
        'output_sha256': sha256_file(output),
        'title': title,
        'copied_body_nodes': copied_counts,
        'removed_body_nodes': removed_counts,
        'normalized_numbering_map': numbering_map,
        'materialized_list_markers': materialized_list_count,
        'pandoc_command': command,
        'pandoc_stdout': completed.stdout.strip(),
        'pandoc_stderr': completed.stderr.strip(),
        'qa': qa,
    }
    report.parent.mkdir(parents=True, exist_ok=True)
    report.write_text(
        json.dumps(receipt, indent=2, ensure_ascii=False) + '\n',
        encoding='utf-8',
    )
    return receipt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description='Stage the revised thesis DOCX while preserving protected regions.'
    )
    parser.add_argument('--reference', required=True)
    parser.add_argument('--source', required=True)
    parser.add_argument('--output', required=True)
    parser.add_argument('--report', required=True)
    parser.add_argument(
        '--pandoc',
        default=r'C:\Users\JM\AppData\Local\Pandoc\pandoc.exe',
    )
    return parser.parse_args()


def main() -> int:
    try:
        receipt = build(parse_args())
    except (SyncError, OSError, ValueError) as exc:
        print(f'ERROR: {exc}')
        return 1
    print(
        'Staged synchronized DOCX:',
        receipt['output'],
        receipt['output_sha256'],
    )
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
