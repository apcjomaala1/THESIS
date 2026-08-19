from __future__ import annotations

import copy
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"C:\Projects\THESIS")
REFERENCE = ROOT / "Finals_Revised_Paper_WASD.docx"
SOURCE = ROOT / "thesis_docs" / "Finals_Revised_Paper_WASD.md"
OUTPUT = ROOT / "Finals_Revised_Paper_WASD_from_markdown.docx"


def clone(value):
    return copy.deepcopy(value) if value is not None else None


def apply_para_template(paragraph, template):
    if template._p.pPr is not None:
        current = paragraph._p.pPr
        if current is not None:
            paragraph._p.remove(current)
        paragraph._p.insert(0, clone(template._p.pPr))


def add_styled_run(paragraph, text, template=None, bold=False, italic=False):
    run = paragraph.add_run(text)
    if template is not None and template._r.rPr is not None:
        if run._r.rPr is not None:
            run._r.remove(run._r.rPr)
        run._r.insert(0, clone(template._r.rPr))
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return run


def add_hyperlink(paragraph, text, url):
    rid = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    props.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)
    run.append(props)
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def math_text(value):
    value = value.strip()
    value = value.replace(r"\(", "").replace(r"\)", "")
    replacements = {
        r"\\mathbf": "", r"\\text": "", r"\\lvert": "|", r"\\rvert": "|",
        r"\\leq": "<=", r"\\geq": ">=", r"\\neq": "!=", r"\\neq": "!=",
        r"\\cap": "∩", r"\\varnothing": "∅", r"\\neq": "≠", r"\\le": "≤",
        r"\\ge": "≥", r"\\times": "×", r"\\cdot": "·", r"\\tau": "τ",
        r"\\delta": "δ", r"\\max": "max", r"\\sum": "sum", r"\\cos": "cos",
    }
    for source, target in replacements.items():
        value = value.replace(source, target)
    value = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", value)
    value = re.sub(r"\\([A-Za-z]+)", r"\1", value)
    value = value.replace("{", "").replace("}", "")
    return re.sub(r"\s+", " ", value).strip()


def add_inline(paragraph, text, template=None):
    pattern = re.compile(r"(\*\*.*?\*\*|\*[^*]+\*|`[^`]+`|\$?\\\(.*?\\\)\$?|\[[^\]]+\]\([^\)]+\))")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            add_styled_run(paragraph, text[cursor:match.start()], template)
        token = match.group(0)
        if token.startswith(("**", "__")):
            add_styled_run(paragraph, token[2:-2], template, bold=True)
        elif token.startswith(("*", "_")):
            add_styled_run(paragraph, token[1:-1], template, italic=True)
        elif token.startswith("`"):
            run = add_styled_run(paragraph, token[1:-1], template)
            run.font.name = "Consolas"
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            add_styled_run(paragraph, math_text(token.strip("$\\()")), template, italic=True)
        cursor = match.end()
    if cursor < len(text):
        add_styled_run(paragraph, text[cursor:], template)


def parse_table(lines):
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def table_widths(column_count):
    if column_count == 2:
        return [Inches(4.6), Inches(1.9)]
    if column_count == 3:
        return [Inches(1.8), Inches(3.65), Inches(1.05)]
    if column_count == 4:
        return [Inches(1.55), Inches(1.55), Inches(1.85), Inches(1.55)]
    return [Inches(6.5 / column_count)] * column_count


def add_table(document, rows, table_template):
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = table_template.style
    table.autofit = False
    widths = table_widths(len(rows[0]))
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            set_cell_width(cell, widths[column_index])
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(math_text(value))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.bold = row_index == 0
    return table


def add_paragraph(document, text, template, style_name=None, alignment=None, italic=False):
    paragraph = document.add_paragraph(style=style_name or template.style.name)
    apply_para_template(paragraph, template)
    if alignment is not None:
        paragraph.alignment = alignment
    add_inline(paragraph, text, template.runs[0] if template.runs else None)
    if italic:
        for run in paragraph.runs:
            run.italic = True
    return paragraph


def main():
    doc = Document(REFERENCE)
    source_paragraphs = doc.paragraphs
    templates = {
        "cover_title": source_paragraphs[1],
        "cover_line": source_paragraphs[2],
        "cover_school": source_paragraphs[4],
        "cover_course": source_paragraphs[5],
        "cover_subject": source_paragraphs[6],
        "cover_by": source_paragraphs[7],
        "cover_author": source_paragraphs[8],
        "heading1": source_paragraphs[12],
        "heading2": source_paragraphs[13],
        "heading3": source_paragraphs[37],
        "body": source_paragraphs[14],
        "caption": source_paragraphs[123],
        "reference": source_paragraphs[183],
        "table": doc.tables[0],
    }
    body = doc._element.body
    for element in list(body):
        if element.tag != qn("w:sectPr"):
            body.remove(element)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    chapter_index = next(i for i, line in enumerate(lines) if line.startswith("# I."))
    cover_lines = [line.strip() for line in lines[:chapter_index] if line.strip()]
    title = re.sub(r"[*_]", "", cover_lines[0])
    cover_map = [
        (title, "cover_title", True),
        (cover_lines[1], "cover_line", False),
        (cover_lines[2], "cover_line", False),
        (cover_lines[3], "cover_school", False),
        (cover_lines[4], "cover_course", False),
        (cover_lines[5], "cover_subject", False),
        (cover_lines[6], "cover_by", False),
    ]
    for value, name, italic in cover_map:
        paragraph = add_paragraph(doc, value, templates[name], italic=italic)
        if name == "cover_title":
            for run in paragraph.runs:
                run.bold = True
                run.italic = True
    for author in cover_lines[7:]:
        add_paragraph(doc, author, templates["cover_author"])
    doc.add_page_break()

    index = chapter_index
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("\\["):
            equation = []
            index += 1
            while index < len(lines) and lines[index].strip() != "\\]":
                equation.append(lines[index].strip())
                index += 1
            paragraph = add_paragraph(
                doc, math_text(" ".join(equation)), templates["body"], alignment=WD_ALIGN_PARAGRAPH.CENTER
            )
            for run in paragraph.runs:
                run.font.name = "Cambria Math"
                run.font.size = Pt(11)
            index += 1
            continue
        if stripped.startswith("|"):
            block = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                block.append(lines[index].strip())
                index += 1
            add_table(doc, parse_table(block), templates["table"])
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            add_paragraph(doc, heading.group(2), templates[f"heading{level}"], f"Heading {level}")
            index += 1
            continue
        caption = re.fullmatch(r"\*(Table .+)\*", stripped)
        if caption:
            add_paragraph(doc, caption.group(1), templates["caption"], alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
            index += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            paragraph = add_paragraph(doc, f"{numbered.group(1)}. {numbered.group(2)}", templates["body"])
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            index += 1
            continue
        template = templates["reference"] if stripped.startswith("[") else templates["body"]
        add_paragraph(doc, stripped, template)
        index += 1

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
