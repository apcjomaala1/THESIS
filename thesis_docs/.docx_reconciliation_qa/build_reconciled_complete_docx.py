from __future__ import annotations

"""Build a complete reconciled thesis without mutating either source DOCX."""

import argparse
import copy
import importlib.util
import json
import re
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


EXPECTED_REFERENCE_SHA256 = (
    '481A95449DD144EC07ADA0DD6E6B90100972BBE287645AE7501DC5606FB582D2'
)
EXPECTED_TEAMMATE_SHA256 = (
    'D1973E367C391AD0A2752804FFD1743CE0B42B4009B0D7C0B57AFECE429F686A'
)
PANDOC_READER = (
    'markdown+pipe_tables+tex_math_dollars+tex_math_single_backslash'
)
CHAPTER_IV = 'IV. RESULTS AND DISCUSSION'
CHAPTER_V = 'V. SUMMARY, CONCLUSIONS, AND RECOMMENDATIONS'
REFERENCES = 'References'


class ReconciliationError(RuntimeError):
    pass


def load_sync_module(path: Path) -> Any:
    spec = importlib.util.spec_from_file_location('verified_docx_sync', path)
    if spec is None or spec.loader is None:
        raise ReconciliationError(f'Cannot load verified synchronizer: {path}')
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip('|').split('|')]


def render_table(rows: list[list[str]], indices: list[int]) -> str:
    projected = [[row[index] for index in indices] for row in rows]
    lines = [
        '| ' + ' | '.join(projected[0]) + ' |',
        '| ' + ' | '.join('---' for _ in indices) + ' |',
    ]
    lines.extend('| ' + ' | '.join(row) + ' |' for row in projected[2:])
    return '\n'.join(lines)


def split_wide_tables(markdown: str) -> str:
    markdown = re.sub(
        r'^### Table 4\.[123]\..*?\n\n',
        '',
        markdown,
        flags=re.MULTILINE,
    )
    lines = markdown.splitlines()
    output: list[str] = []
    index = 0
    while index < len(lines):
        if (
            lines[index].lstrip().startswith('|')
            and index + 1 < len(lines)
            and re.match(r'^\s*\|?\s*:?-+', lines[index + 1])
        ):
            block: list[str] = []
            while index < len(lines) and lines[index].lstrip().startswith('|'):
                block.append(lines[index])
                index += 1
            rows = [parse_table_row(line) for line in block]
            width = len(rows[0])
            first = rows[0][0]
            if first == 'Method' and width == 12:
                output.extend([
                    '**Table 4.1a. Ranking-Metric Performance**',
                    '',
                    render_table(rows, [0, 1, 2, 3, 4]),
                    '',
                    '**Table 4.1b. Thresholded Operating-Point Performance**',
                    '',
                    render_table(rows, [0, 5, 6, 7]),
                    '',
                    '**Table 4.1c. Confusion Counts at Frozen Thresholds**',
                    '',
                    render_table(rows, [0, 8, 9, 10, 11]),
                ])
            elif first == 'Comparison' and width == 6:
                output.extend([
                    '**Table 4.2a. Paired Differences in PR-AUC and F0.5**',
                    '',
                    render_table(rows, [0, 1, 2, 5]),
                    '',
                    '**Table 4.2b. Paired Differences in Precision and Recall**',
                    '',
                    render_table(rows, [0, 3, 4, 5]),
                ])
            elif first == 'Method' and width == 9:
                output.extend([
                    '**Table 4.3a. Validation and Test PR-AUC and F0.5**',
                    '',
                    render_table(rows, [0, 1, 2, 3, 4]),
                    '',
                    '**Table 4.3b. Validation and Test Precision and Recall**',
                    '',
                    render_table(rows, [0, 5, 6, 7, 8]),
                ])
            else:
                output.extend(block)
            continue
        output.append(lines[index])
        index += 1
    return '\n'.join(output).strip() + '\n'


def word_markdown(chapter_iv: Path, chapter_v: Path) -> str:
    text = (
        chapter_iv.read_text(encoding='utf-8').strip()
        + '\n\n'
        + chapter_v.read_text(encoding='utf-8').strip()
        + '\n\n# References\n'
    )
    text = split_wide_tables(text)
    text = re.sub(r'(?m)^(\d+)\.\s+', r'\1\\. ', text)
    separated: list[str] = []
    for line in text.splitlines():
        separated.append(line)
        if re.match(r'^\d+\\\.\s', line):
            separated.append('')
    text = '\n'.join(separated) + '\n'
    replacements = {
        '\u2013': '-',
        '\u2014': '-',
        '\u2018': "'",
        '\u2019': "'",
        '\u201c': '"',
        '\u201d': '"',
        '\u00a0': ' ',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def page_break_clone(sync: Any, reference_break: Any) -> Any:
    cloned = copy.deepcopy(reference_break)
    if not sync.is_page_break(cloned):
        raise ReconciliationError('Reference page-break clone is invalid')
    return cloned


def heading_clone(sync: Any, document: Any, text: str) -> Any:
    _, heading = sync.find_anchor(document, text)
    cloned = copy.deepcopy(heading)
    for tag in (sync.W_BOOKMARK_START, sync.W_BOOKMARK_END):
        for bookmark in list(cloned.iter(tag)):
            parent = bookmark.getparent()
            if parent is not None:
                parent.remove(bookmark)
    return cloned


def insert_before(document: Any, anchor: Any, nodes: list[Any]) -> None:
    body = document.element.body
    insertion = list(body.iterchildren()).index(anchor)
    for node in nodes:
        body.insert(insertion, node)
        insertion += 1


def chapter_fingerprint(sync: Any, document: Any, heading: str) -> str:
    children = sync.direct_children(document)
    _, anchor = sync.find_anchor(document, heading)
    start = children.index(anchor)
    end = next(
        index
        for index in range(start + 1, len(children))
        if sync.is_page_break(children[index])
    )
    return sync.canonical_fingerprint(children[start:end])


def set_cell_margins(cell: Any, top: int = 80, start: int = 100,
                     bottom: int = 80, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn('w:tcMar'))
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for edge, value in (
        ('top', top), ('start', start), ('bottom', bottom), ('end', end)
    ):
        node = tc_mar.find(qn(f'w:{edge}'))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def polish_table(sync: Any, table: Any, widths: list[int],
                 font_size: float) -> None:
    sync.set_content_table_geometry(table, widths)
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if column_index == 0
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    if row_index == 0:
                        run.bold = True


def set_visible_list_indents(sync: Any, roots: list[Any]) -> int:
    count = 0
    for root in roots:
        for paragraph in root.iter(sync.W_P):
            text = sync.node_text(paragraph)
            if not re.match(r'^\d+\.\s', text):
                continue
            p_pr = paragraph.find(sync.W_P_PR)
            if p_pr is None:
                p_pr = OxmlElement('w:pPr')
                paragraph.insert(0, p_pr)
            indent = p_pr.find(qn('w:ind'))
            if indent is None:
                indent = OxmlElement('w:ind')
                p_pr.append(indent)
            indent.set(qn('w:left'), '720')
            indent.set(qn('w:hanging'), '360')
            count += 1
    return count


def count_visible_list_paragraphs(sync: Any, document: Any) -> int:
    return sum(
        bool(re.match(r'^\d+\.\s', sync.node_text(paragraph)))
        for paragraph in document.element.body.iter(sync.W_P)
    )


def table_geometry_ok(table: Any, widths: list[int]) -> bool:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    grid = table._tbl.tblGrid
    grid_widths = [
        int(node.get(qn('w:w'), '0'))
        for node in grid.findall(qn('w:gridCol'))
    ]
    if tbl_w is None or int(tbl_w.get(qn('w:w'), '0')) != sum(widths):
        return False
    if grid_widths != widths:
        return False
    for row in table.rows:
        actual = [
            int(cell._tc.get_or_add_tcPr().find(qn('w:tcW')).get(qn('w:w')))
            for cell in row.cells
        ]
        if actual != widths:
            return False
    return True


def document_text(document: Any) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return '\n'.join(parts)


def build(args: argparse.Namespace) -> dict[str, Any]:
    reference = Path(args.reference).resolve()
    teammate = Path(args.teammate).resolve()
    chapter_iv = Path(args.chapter_iv).resolve()
    chapter_v = Path(args.chapter_v).resolve()
    synchronizer = Path(args.synchronizer).resolve()
    pandoc = Path(args.pandoc).resolve()
    output = Path(args.output).resolve()
    report = Path(args.report).resolve()
    for path in (
        reference, teammate, chapter_iv, chapter_v, synchronizer, pandoc
    ):
        if not path.exists():
            raise ReconciliationError(f'Required input does not exist: {path}')
    if output in {reference, teammate}:
        raise ReconciliationError('Output must not overwrite either source DOCX')
    if output.exists():
        raise ReconciliationError(f'Refusing to overwrite output: {output}')

    sync = load_sync_module(synchronizer)
    reference_hash = sync.sha256_file(reference)
    teammate_hash = sync.sha256_file(teammate)
    if reference_hash != EXPECTED_REFERENCE_SHA256:
        raise ReconciliationError(
            f'Unexpected reference hash: {reference_hash}'
        )
    if teammate_hash != EXPECTED_TEAMMATE_SHA256:
        raise ReconciliationError(
            f'Unexpected teammate hash: {teammate_hash}'
        )

    reference_doc = Document(reference)
    reference_regions = sync.reference_regions(reference_doc)
    protected_before = {
        'cover_without_title': sync.canonical_fingerprint(
            reference_regions['cover_without_title']
        ),
        'cover_author_table': sync.canonical_fingerprint(
            [reference_regions['author_table']]
        ),
        'references': sync.canonical_fingerprint(
            reference_regions['references']
        ),
        'chapter_i': chapter_fingerprint(sync, reference_doc, 'I. Introduction'),
        'chapter_ii': chapter_fingerprint(
            sync, reference_doc, 'II. RELATED WORK'
        ),
        'chapter_iii': chapter_fingerprint(
            sync, reference_doc, 'III. METHODOLOGY'
        ),
    }

    markdown = word_markdown(chapter_iv, chapter_v)
    with tempfile.TemporaryDirectory(prefix='thesis-reconcile-') as temp_dir:
        temp_root = Path(temp_dir)
        markdown_path = temp_root / 'chapters_iv_v.md'
        generated_path = temp_root / 'chapters_iv_v.docx'
        markdown_path.write_text(markdown, encoding='utf-8')
        command = [
            str(pandoc),
            str(markdown_path),
            '-f', PANDOC_READER,
            '-t', 'docx',
            '--reference-doc', str(reference),
            '-o', str(generated_path),
        ]
        completed = subprocess.run(
            command,
            check=False,
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='replace',
        )
        if completed.returncode != 0 or not generated_path.exists():
            raise ReconciliationError(
                'Pandoc generation failed: '
                + (completed.stderr.strip() or completed.stdout.strip())
            )
        generated = Document(generated_path)
        chapter_iv_nodes = sync.source_inner_nodes(
            generated, CHAPTER_IV, CHAPTER_V
        )
        chapter_v_nodes = sync.source_inner_nodes(
            generated, CHAPTER_V, REFERENCES
        )

        output_doc = Document(reference)
        _, refs_anchor = sync.find_anchor(output_doc, REFERENCES)
        refs_break = sync.page_break_before(output_doc, refs_anchor)
        nodes = (
            [
                page_break_clone(sync, refs_break),
                heading_clone(sync, generated, CHAPTER_IV),
            ]
            + chapter_iv_nodes
            + [
                page_break_clone(sync, refs_break),
                heading_clone(sync, generated, CHAPTER_V),
            ]
            + chapter_v_nodes
        )
        insert_before(output_doc, refs_break, nodes)

        new_tables = output_doc.tables[3:]
        widths = [
            [1500, 2700, 1720, 1720, 1720],
            [2700, 2220, 2220, 2220],
            [3000, 1590, 1590, 1590, 1590],
            [2700, 2100, 2100, 2460],
            [2700, 2100, 2100, 2460],
            [2500, 1715, 1715, 1715, 1715],
            [2500, 1715, 1715, 1715, 1715],
        ]
        if len(new_tables) != len(widths):
            raise ReconciliationError(
                f'Expected {len(widths)} new tables; found {len(new_tables)}'
            )
        for table, table_widths in zip(new_tables, widths):
            polish_table(
                sync,
                table,
                table_widths,
                7.5 if len(table_widths) == 5 else 8.0,
            )
        new_list_count = set_visible_list_indents(sync, nodes)
        if new_list_count != 25:
            raise ReconciliationError(
                f'Expected 25 new visible list paragraphs; found {new_list_count}'
            )
        output.parent.mkdir(parents=True, exist_ok=True)
        output_doc.save(output)

    reopened = Document(output)
    output_regions = sync.reference_regions(reopened)
    protected_after = {
        'cover_without_title': sync.canonical_fingerprint(
            output_regions['cover_without_title']
        ),
        'cover_author_table': sync.canonical_fingerprint(
            [output_regions['author_table']]
        ),
        'references': sync.canonical_fingerprint(output_regions['references']),
        'chapter_i': chapter_fingerprint(sync, reopened, 'I. Introduction'),
        'chapter_ii': chapter_fingerprint(sync, reopened, 'II. RELATED WORK'),
        'chapter_iii': chapter_fingerprint(sync, reopened, 'III. METHODOLOGY'),
    }
    if protected_after != protected_before:
        raise ReconciliationError(
            'Protected cover, Chapters I-III, or References changed'
        )

    headings = sync.heading_counts(reopened)
    expected_headings = {'Heading 1': 6, 'Heading 2': 26, 'Heading 3': 20}
    if headings != expected_headings:
        raise ReconciliationError(f'Unexpected heading counts: {headings}')
    if sync.count_direct_page_breaks(reopened) != 5:
        raise ReconciliationError('Expected five direct page breaks')
    if sync.count_hyperlinks(reopened) != 13:
        raise ReconciliationError('Expected 13 preserved hyperlinks')
    if sync.count_numbered_paragraphs(reopened) != 0:
        raise ReconciliationError('Automatic numbering remains in output')
    list_count = count_visible_list_paragraphs(sync, reopened)
    # The preserved Chapters I-III contain 18 numbered paragraphs plus one
    # bullet; Chapters IV-V add 25 numbered paragraphs.
    expected_lists = 43
    if list_count != expected_lists:
        raise ReconciliationError(
            f'Expected {expected_lists} visible list paragraphs; found {list_count}'
        )
    special = sync.tracked_or_field_counts(reopened)
    if any(special.values()):
        raise ReconciliationError(
            f'Unexpected tracked changes/comments/fields: {special}'
        )
    if len(reopened.sections) != 1:
        raise ReconciliationError(
            f'Expected one section; found {len(reopened.sections)}'
        )
    geometry_checks = [
        table_geometry_ok(table, table_widths)
        for table, table_widths in zip(reopened.tables[3:], widths)
    ]
    if not all(geometry_checks):
        raise ReconciliationError('One or more result-table geometries failed')

    text = document_text(reopened)
    stale_patterns = [
        '22,798', 'score_ewma', 'risk_spike', 'risk_drop',
        'p < 0.05', 'Equivalent, CI spans 0', 'over 99%',
        '100 times faster', '< 1 ms', 'eliminates over 95%',
    ]
    stale_found = [pattern for pattern in stale_patterns if pattern in text]
    if stale_found:
        raise ReconciliationError(
            f'Stale teammate claims remain: {stale_found}'
        )
    if '**' in text:
        raise ReconciliationError('Literal Markdown emphasis leaked into DOCX')

    table_shapes = [
        [len(table.rows), len(table.columns)]
        for table in reopened.tables
    ]
    receipt = {
        'schema_version': 1,
        'created_utc': datetime.now(timezone.utc).isoformat(),
        'reference': str(reference),
        'teammate_source': str(teammate),
        'chapter_iv_source': str(chapter_iv),
        'chapter_v_source': str(chapter_v),
        'output': str(output),
        'reference_sha256': reference_hash,
        'teammate_sha256': teammate_hash,
        'chapter_iv_sha256': sync.sha256_file(chapter_iv),
        'chapter_v_sha256': sync.sha256_file(chapter_v),
        'output_sha256': sync.sha256_file(output),
        'pandoc_command': command,
        'qa': {
            'protected_fingerprints': protected_after,
            'heading_counts': headings,
            'direct_page_breaks': 5,
            'hyperlinks': 13,
            'visible_list_paragraphs': list_count,
            'automatic_numbered_paragraphs': 0,
            'special_content_counts': special,
            'table_shapes': table_shapes,
            'result_table_geometry_checks': geometry_checks,
            'stale_claim_patterns_found': stale_found,
            'sections': len(reopened.sections),
        },
        'experiment_boundary': {
            'retrained': False,
            'retuned': False,
            'final_test_rescored': False,
        },
    }
    report.parent.mkdir(parents=True, exist_ok=True)
    report.write_text(
        json.dumps(receipt, indent=2, ensure_ascii=False) + '\n',
        encoding='utf-8',
    )
    return receipt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument('--reference', required=True)
    parser.add_argument('--teammate', required=True)
    parser.add_argument('--chapter-iv', required=True)
    parser.add_argument('--chapter-v', required=True)
    parser.add_argument('--synchronizer', required=True)
    parser.add_argument('--output', required=True)
    parser.add_argument('--report', required=True)
    parser.add_argument(
        '--pandoc',
        default=r'C:\Users\JM\AppData\Local\Pandoc\pandoc.exe',
    )
    return parser.parse_args()


def main() -> int:
    try:
        receipt = build(parse_args())
    except (OSError, ValueError, ReconciliationError) as exc:
        print(f'ERROR: {exc}')
        return 1
    print('Built reconciled DOCX:', receipt['output'])
    print('SHA-256:', receipt['output_sha256'])
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
